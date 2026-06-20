import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('软件体系结构与设计模式 - 课程设计报告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('\n')
    
    # 1. 选择主题
    doc.add_heading('1. 选择主题', level=1)
    doc.add_paragraph('本项目选择的主题为《学业发展规划系统 (Qt桌面版)》。本项目旨在帮助学生全面管理个人学业、项目经验、角色活动和职业目标，并通过集成大语言模型（LLM）提供智能化的简历生成和学业规划建议，实现对学生发展的全生命周期管理。')
    
    # 2. 体系结构风格
    doc.add_heading('2. 使用的体系结构风格', level=1)
    doc.add_paragraph('本项目采用了以下混合体系结构风格：')
    doc.add_paragraph('1. 分层架构模式 (Layered Architecture)：后端按照表现层（API/Client）、业务逻辑层（Service）、数据访问层（DAO）、数据模型层（Model）进行严格分层，保证高内聚和低耦合。', style='List Bullet')
    doc.add_paragraph('2. 客户端-服务器模式 (Client-Server)：系统分为 C++ Qt 客户端以及本地运行的 Flask AI 服务器，通过 HTTP RESTful API 进行跨进程通信。', style='List Bullet')
    doc.add_paragraph('3. 协调器模式 (Coordinator Pattern)：在客户端内部，采用 AppShellController 和 DataRefreshCoordinator 等中心协调器管理页面状态、跨组件通信和数据全局刷新。', style='List Bullet')
    
    # 3. 需求说明
    doc.add_heading('3. 需求说明', level=1)
    doc.add_paragraph('系统包含以下核心需求：')
    doc.add_paragraph('1. 学业信息管理：支持课程、成绩、角色、活动的 CRUD（增删改查）操作。', style='List Bullet')
    doc.add_paragraph('2. 目标与时间轴跟踪：记录学期目标完成度，并以时间轴形式可视化展示发展轨迹。', style='List Bullet')
    doc.add_paragraph('3. 简历生成与管理：通过组合管理的内容素材，动态拼接生成正式的个人简历。', style='List Bullet')
    doc.add_paragraph('4. AI智能辅助：提供基于本地大模型的学业规划分析、简历优化和悬浮菜单划词提问功能。', style='List Bullet')
    
    # 4. 用例分析
    doc.add_heading('4. 用例分析', level=1)
    doc.add_paragraph('核心用例：AI 辅助简历生成与优化')
    doc.add_paragraph('参与者：学生用户')
    doc.add_paragraph('前置条件：用户已在系统中录入项目经历和活动。')
    doc.add_paragraph('主事件流：\n1. 用户在“项目经验”页面选中一段描述文本。\n2. 用户释放鼠标，系统弹出“作为提示词上传”悬浮菜单。\n3. 用户点击后，文本被传递给 AI 面板。\n4. AI 服务分析文本并生成适合简历的正式表述。\n5. 用户点击“应用到简历”，系统将内容回填至简历生成页。')
    
    # 5. 静态设计
    doc.add_heading('5. 静态设计 (类与模块)', level=1)
    doc.add_paragraph('系统的核心类图和模块划分为：')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '层级'
    hdr_cells[1].text = '核心类'
    hdr_cells[2].text = '职责说明'
    
    data = [
        ('UI 层', 'MainWindow, BasePage, SidebarWidget', '展示界面，接收用户交互'),
        ('协调层', 'AppShellController, AiContextMediator', '管理组件通信、全局状态刷新'),
        ('Service 层', 'CourseService, AiService', '封装业务规则、调用外部模型 API'),
        ('DAO 层', 'CourseDao, DaoBase', '处理 SQLite 数据库的持久化操作'),
        ('Model 层', 'Course, Goal, Resume', '业务实体定义和数据承载')
    ]
    for layer, classes, duty in data:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = classes
        row_cells[2].text = duty

    # 6. 详细设计与分工
    doc.add_heading('6. 详细设计与团队分工', level=1)
    doc.add_paragraph('根据技术解析文档，团队分工如下：')
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '团队成员'
    hdr_cells2[1].text = '负责模块 (详细设计)'
    hdr_cells2[2].text = '核心开发文件'
    
    data2 = [
        ('组长', '架构设计、客户端壳层开发、核心协调器与AI交互设计', 'HttpServer.h, MainWindow.cpp, core/*'),
        ('组员A', '后端API层开发、路由设计、JSON序列化工具', 'api/*, util/JsonUtils.h'),
        ('组员B', '核心Model设计、DAO持久化层实现、Service业务逻辑', 'model/*, dao/*, service/*'),
        ('组员C', '测试用例编写、系统测试、相关技术文档编制', 'tests/*, docs/*')
    ]
    for member, module, files in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = member
        row_cells[1].text = module
        row_cells[2].text = files

    # 7. 时序图与核心功能状态
    doc.add_heading('7. 时序交互与核心功能状态', level=1)
    doc.add_paragraph('AI 建议生成时序：\n用户(UI) -> AiContextMediator(中介) -> AiPanelWidget(UI面板) -> AiService(业务) -> 本地Qwen模型(HTTP) -> AiService返回分析结果 -> UI 更新显示。')
    doc.add_paragraph('核心功能状态转移 (目标管理生命周期)：\n[草稿 Draft] -> (启动) -> [进行中 In Progress] -> (达成) -> [已完成 Completed] / (放弃) -> [已取消 Cancelled]。')

    # 8. GitHub 项目链接与阶段规划
    doc.add_heading('8. GitHub 提交规划与项目链接', level=1)
    doc.add_paragraph('项目代码托管平台：Github')
    doc.add_paragraph('项目链接：https://github.com/xxxx/pdp-desktop (待替换)')
    doc.add_paragraph('\n为了体现合理的敏捷开发发展顺序，项目代码向 GitHub 上传分为 6 个 Stage：')
    
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'Stage (阶段)'
    hdr_cells3[1].text = '里程碑目标'
    hdr_cells3[2].text = '包含的主要内容'
    
    data3 = [
        ('Stage 1', '项目初始化与架构搭建', '基础 Qt 项目框架、CMake 构建配置、数据库 Schema 初始化'),
        ('Stage 2', '核心数据与持久化层', '集成 SQLite，开发基础的数据模型 (Models) 和 DAO 类'),
        ('Stage 3', '业务逻辑与API层封装', 'Service 业务规则实现，QHttpServer 路由映射开发'),
        ('Stage 4', '桌面端基础框架与导航', '主窗体壳层、侧边栏导航、统一样式表和基类页面'),
        ('Stage 5', '详细界面与CRUD交互', '完善各大页面的表单组件、图表分析与数据刷新协调机制'),
        ('Stage 6', 'AI集成、简历引擎与收尾', '集成 Python AI 服务调用、实现简历生成和整体UI视觉打磨')
    ]
    for stage, goal, content in data3:
        row_cells = table3.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = goal
        row_cells[2].text = content

    # 9. 进度表 (甘特图规划)
    doc.add_heading('9. 项目进度表规划', level=1)
    doc.add_paragraph('以下为项目开发的 6 周进度安排（以周为单位的甘特图表征）：')
    
    table4 = doc.add_table(rows=7, cols=7)
    table4.style = 'Table Grid'
    headers = ['任务名称', 'W1', 'W2', 'W3', 'W4', 'W5', 'W6']
    for i, h in enumerate(headers):
        table4.cell(0, i).text = h
        
    gantt_data = [
        ('Stage 1: 架构设计与环境搭建', '██', '', '', '', '', ''),
        ('Stage 2: DAO与数据库开发', '', '██', '', '', '', ''),
        ('Stage 3: Service与API开发', '', '██', '██', '', '', ''),
        ('Stage 4: 前端核心壳层开发', '', '', '██', '██', '', ''),
        ('Stage 5: 业务组件与数据绑定', '', '', '', '██', '██', ''),
        ('Stage 6: AI集成与测试联调', '', '', '', '', '██', '██')
    ]
    
    for row_idx, row_data in enumerate(gantt_data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            table4.cell(row_idx, col_idx).text = cell_data

    doc.save('软件体系结构与设计模式-设计报告.docx')
    print('Report generated successfully!')

if __name__ == '__main__':
    create_report()
