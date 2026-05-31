import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_presentation_doc():
    doc = Document()
    
    # 标题
    title = doc.add_heading('《学业发展规划系统(Qt桌面版)》期末演示总控文档与演讲稿', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # ==========================================
    # 第一部分：总体演示规划
    # ==========================================
    doc.add_heading('第一部分：总体演示规划 (总时长：10分钟)', level=1)
    
    doc.add_heading('1. 演示内容与时间分配表', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '主讲人'
    hdr_cells[1].text = '时长'
    hdr_cells[2].text = '负责内容模块 (对应PPT与要求)'
    hdr_cells[3].text = '现场操作配合'
    
    data = [
        ('组长', '2.5分钟', '项目背景、组织分工、软件体系结构设计（架构图展示）、系统扩展性体现', '播放PPT：展示项目背景与整体C/S架构、分层架构图'),
        ('组员A', '2分钟', 'API层与路由设计、JSON序列化工具、系统复用性体现（前后端解耦）', '播放PPT：讲解路由分发机制；配合打开Qt客户端展示网络加载状态'),
        ('组员B', '3分钟', 'Model/DAO/Service层逻辑、核心功能与创新功能演示（重点演示AI辅助简历生成）', '实操软件：录入一段活动经历，选中文字呼出“AI悬浮菜单”，演示AI改写并生成简历'),
        ('组员C', '2.5分钟', '软件质量意识（异常处理、边界情况）、可维护性体现、项目亮点与总结', '实操软件：演示必填项为空的校验提示、关闭大模型服务后系统的优雅降级（错误捕获）')
    ]
    for speaker, time, content, action in data:
        row_cells = table.add_row().cells
        row_cells[0].text = speaker
        row_cells[1].text = time
        row_cells[2].text = content
        row_cells[3].text = action

    # ==========================================
    # 第二部分：各成员演讲思路与逐字稿
    # ==========================================
    doc.add_heading('第二部分：个人演讲思路与逐字稿', level=1)
    
    # --- 组长 ---
    doc.add_heading('【组长】开场、背景与体系结构设计 (2.5分钟)', level=2)
    p = doc.add_paragraph()
    p.add_run('【演讲思路】\n').bold = True
    p.add_run('作为团队Leader，负责把控全场节奏。首先交代为什么要做这个项目，引出小组成员的明确分工。接着切入硬核的“体系结构设计”，重点讲解 C/S 混合架构和 Qt 协调器模式（Coordinator Pattern），并回答“扩展性体现在哪里”这一考核点。')
    p = doc.add_paragraph()
    p.add_run('【逐字稿】\n').bold = True
    p.add_run('各位老师、同学大家好。我们组带来的是《学业发展规划系统 Qt 桌面版》。\n')
    p.add_run('首先介绍【项目背景】：大学生在校期间往往面临学业繁杂、目标不清晰、简历素材零散等痛点。为了实现全生命周期的学业管理，我们团队分工协作：我负责整体架构与核心UI，组员A主攻API层，组员B开发模型与底层DAO，组员C负责系统测试与质量把控。\n')
    p.add_run('【切换PPT至架构图】在【体系结构设计】上，我们采用了经典的 C/S 混合架构与严格的 MVC 分层模式。大家可以看到，底层通过 SQLite 嵌入式存储，向上依次是 DAO 抽象层、Service 业务逻辑层，以及由 QHttpServer 驱动的 API 层。前端则是纯 Qt 构建的复杂 GUI。\n')
    p.add_run('在这个架构下，我们项目的【扩展性与复用性】得到了极大的体现。我们在前端引入了“核心协调器模式（Coordinator Pattern）”，例如 DataRefreshCoordinator 和 AiContextMediator，它们把页面刷新和组件通信完全解耦了。哪怕我们未来再加十个新页面，底层数据流都不需要做任何修改逻辑。下面有请组员A为大家详细讲解 API 层的解耦设计。')
    
    # --- 组员A ---
    doc.add_heading('【组员A】API路由设计与系统复用性 (2分钟)', level=2)
    p = doc.add_paragraph()
    p.add_run('【演讲思路】\n').bold = True
    p.add_run('紧跟组长的架构图，聚焦在“前后端通信的桥梁”。重点介绍路由的设计、跨域处理，并解释 JSON 序列化工具的复用性。')
    p = doc.add_paragraph()
    p.add_run('【逐字稿】\n').bold = True
    p.add_run('大家好，我是组员A，主要负责后端的 API 层与路由开发。\n')
    p.add_run('正如组长提到的高扩展性，我们在 Qt 内部集成了一个轻量级的 HttpServer。在这个层面上，我们高度体现了系统的【复用性】。我们统一封装了 `JsonUtils`，所有的接口返回都遵循标准化格式（包含 code、message 和 data），这使得前端处理网络请求的代码可以高度复用。\n')
    p.add_run('另外，我们的路由采用了完全的 RESTful 风格（如 GET/POST/PUT/DELETE 分离）。这种设计的最大优势是：虽然我们现在是用 Qt 开发桌面端，但我们的核心业务 API 逻辑可以直接复用于任何 Web 前端（如 Vue）或移动端小程序，实现了前后端的彻底解耦。接下来，组员B将为大家演示这些数据是如何在底层持久化，以及我们的核心亮点功能。')
    
    # --- 组员B ---
    doc.add_heading('【组员B】底层设计与核心/创新功能演示 (3分钟)', level=2)
    p = doc.add_paragraph()
    p.add_run('【演讲思路】\n').bold = True
    p.add_run('简要带过 DAO 与 Service 层的低耦合设计后，将大部分时间留给“现场操作软件”。重点演示项目最大的创新点：基于大模型的 AI 悬浮菜单提取经历并生成简历。')
    p.add_run('【现场操作提示】\n').italic = True
    p.add_run('打开软件 -> 进入“项目经历”页 -> 鼠标高亮选中一段普通文字 -> 弹出“作为提示词上传” -> 点击 -> 侧边栏AI生成专业建议 -> 点击“应用到简历” -> 切换到“简历页”展示成果。')
    p = doc.add_paragraph()
    p.add_run('【逐字稿】\n').bold = True
    p.add_run('老师好，我是组员B。在数据访问层，我们利用面向对象的多态性设计了 DAO 基类，极大地提升了系统的【可维护性】。\n')
    p.add_run('现在请大家看屏幕，我将为大家演示本系统的【核心创新功能——AI 智能学业助手】。传统的系统只能死板地记录数据，而我们接入了本地部署的 Qwen 大语言模型。\n')
    p.add_run('【操作软件演示】例如，我刚刚录入了一段非常口语化的学生会活动经历。现在，我只需用鼠标框选这段文本，系统会通过底层事件拦截，在光标处智能弹出一个“作为提示词上传”的悬浮按钮。点击后，AI 面板会立刻分析这段文本，并将其转换为符合 STAR 法则的专业简历话术。\n')
    p.add_run('随后，我们可以一键将这条由 AI 优化过的内容回填并同步到“简历生成”模块中。这不仅仅是一个 CRUD 系统，而是一个能真正帮助学生打磨求职竞争力的智能工作台。最后，有请组员C为大家展示我们的质量测试与异常处理机制。')
    
    # --- 组员C ---
    doc.add_heading('【组员C】软件质量保障、异常处理与总结 (2.5分钟)', level=2)
    p = doc.add_paragraph()
    p.add_run('【演讲思路】\n').bold = True
    p.add_run('切中“考核要求”中关于软件质量、输入校验和边界情况的要求。通过实操演示软件在非理想状态下的表现，证明系统的健壮性。最后做全组总结。')
    p.add_run('【现场操作提示】\n').italic = True
    p.add_run('1. 尝试创建一个名字为空的目标，展示边界校验。2. (提前在后台关掉AI服务脚本)，点击AI面板请求分析，展示系统的超时捕获和优雅报错提示，而不是崩溃。')
    p = doc.add_paragraph()
    p.add_run('【逐字稿】\n').bold = True
    p.add_run('老师好，我是组员C。一个优秀的软件不仅要有亮眼的功能，还要有极高的稳定性。在开发过程中，我们始终秉持着【软件质量意识】。\n')
    p.add_run('【操作演示1：边界校验】在所有的数据录入模块（如目标制定），我们都加入了严密的输入校验。如果用户试图提交空白数据或非法的日期范围，系统会在前端拦截并弹出规范的 Toast 提示，避免脏数据进入数据库。\n')
    p.add_run('【操作演示2：异常处理机制】更重要的是，我们考虑了各种边界情况。比如现在，我故意断开了后端的本地 AI 服务，模拟网络故障或模型未加载完成。此时我们点击分析，可以看到软件并没有卡死或崩溃，而是利用我们在 QNetworkAccessManager 中专门设置的 `kAiGenerationTimeoutMs` 异步超时机制，捕获了异常，并友善地提示用户“大模型服务当前不可用”，实现了系统的优雅降级。\n')
    p.add_run('【总结】总而言之，我们的项目不仅实现了完整的个人规划闭环，使用了解耦清晰的分层体系结构，更通过集成 AI 大模型赋予了其创新价值，并在异常控制上做到了高容错率。以上就是我们组的全部演示，感谢老师和同学们的聆听，欢迎提问！')
    
    doc.save('期末演示演讲稿与整体规划.docx')
    print('Presentation document generated successfully!')

if __name__ == '__main__':
    create_presentation_doc()
